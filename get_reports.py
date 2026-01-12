import pandas as pd
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor
import heapq
from tqdm import tqdm

CSV_DIR = Path("players_csv")
REPORT_DIR = Path("reports")
REPORT_DIR.mkdir(exist_ok=True)

YEARS = [2022, 2023, 2024, 2025]
HIGHLIGHT_YEARS = [2022, 2023, 2024]

def count_last_zeros(player_csv):
    df = pd.read_csv(player_csv)
    if df.empty:
        return None, 0, {}

    # ✅ même logique que le code debug
    team = df['teamAbbrev'].tolist()[0]

    year_zeros = {}

    for year in YEARS:
        filtered = df[df['season'] == year]

        # ✅ plus d'inversion
        goals_list = filtered['goals'].tolist()

        zeros_list = []
        for idx, g in enumerate(goals_list):
            if g != 0:
                counter = 0
                next_idx = idx + 1
                while next_idx < len(goals_list) and goals_list[next_idx] == 0:
                    counter += 1
                    next_idx += 1
                zeros_list.append(counter)

        # ✅ même pop que dans le debug
        if zeros_list:
            zeros_list.pop()

        year_zeros[year] = zeros_list

    # ✅ last_zero_count compté depuis le match le plus récent
    all_goals = df["goals"].tolist()

    last_zero_count = 0
    for g in all_goals:
        if g == 0:
            last_zero_count += 1
        else:
            break

    return team, last_zero_count, year_zeros


def highlight_paragraph(paragraph, condition):
    if condition:
        for run in paragraph.runs:
            run.font.color.rgb = RGBColor(255, 0, 0)
            run.bold = True

def top3(lst):
    return heapq.nlargest(3, lst) if lst else []

def generate_word_0(csv_folder, output_path):
    if output_path.exists():
        output_path.unlink()
    doc = Document()
    doc.add_heading("Players Zeros by Teams", level=1)

    csv_files = list(csv_folder.glob("*.csv"))
    players_info = []

    for file in tqdm(csv_files, desc="Analyse joueurs"):
        team, last_zero_count, year_zeros = count_last_zeros(file)
        if team is None:
            continue
        players_info.append((team, file.stem, last_zero_count, year_zeros))

    players_info.sort(key=lambda x: x[0])
    current_team = None
    for team, player_name, last_zero_count, year_zeros in players_info:
        if team != current_team:
            current_team = team
            doc.add_heading(f"Team: {team}", level=2)
        p = doc.add_paragraph(f"{player_name}: {last_zero_count}")
        p.runs[0].font.size = Pt(13)
        for year in YEARS:
            zeros_list = year_zeros.get(year, [])
            perc = zeros_list.count(last_zero_count)/len(zeros_list) if zeros_list else 0
            para = doc.add_paragraph(f"{year}: {perc:.2f}", style='List Bullet')
            if perc >= 0.2:
                highlight_paragraph(para, True)
    doc.save(output_path)
    print(f"✅ Word généré : {output_path.name}")

def generate_game_without_goals(csv_folder, output_path):
    if output_path.exists():
        output_path.unlink()
    doc = Document()
    doc.add_heading("Matchs sans marquer", level=1)

    csv_files = list(csv_folder.glob("*.csv"))
    players_info = []

    for file in tqdm(csv_files, desc="Analyse joueurs"):
        team, last_zero_count, year_zeros = count_last_zeros(file)
        if team is None or last_zero_count == 0:
            continue
        players_info.append((team, file.stem, last_zero_count, year_zeros))

    players_info.sort(key=lambda x: x[0])
    current_team = None
    for team, player_name, last_zero_count, year_zeros in players_info:
        if team != current_team:
            current_team = team
            doc.add_heading(f"Team: {team}", level=2)
        p = doc.add_paragraph(f"{player_name}: {last_zero_count}")
        p.runs[0].font.size = Pt(13)
        for year in YEARS:
            top3_list = top3(year_zeros.get(year, []))
            doc.add_paragraph(f"• {year}: {top3_list}")
        combined_top3 = [val for y in HIGHLIGHT_YEARS for val in top3(year_zeros.get(y, []))]
        highlight_paragraph(p, any(last_zero_count >= val for val in combined_top3))

    doc.save(output_path)
    print(f"✅ Word généré : {output_path.name}")

# Exécution
if __name__ == "__main__":
    generate_word_0(CSV_DIR, REPORT_DIR / "0_joueur.docx")
    generate_game_without_goals(CSV_DIR, REPORT_DIR / "game_without_goal.docx")
